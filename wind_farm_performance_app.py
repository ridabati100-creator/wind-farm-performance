import streamlit as st
import sqlite3
import json
import io
from datetime import datetime, date, timedelta

import pandas as pd
import plotly.graph_objects as go
from docx import Document
from docx.shared import Pt, RGBColor

# ============================================================
# Constants
# ============================================================

DAYS = [
    {"key": "mon", "label": "Monday", "short": "Mon"},
    {"key": "tue", "label": "Tuesday", "short": "Tue"},
    {"key": "wed", "label": "Wednesday", "short": "Wed"},
    {"key": "thu", "label": "Thursday", "short": "Thu"},
    {"key": "fri", "label": "Friday", "short": "Fri"},
    {"key": "sat", "label": "Saturday", "short": "Sat"},
    {"key": "sun", "label": "Sunday", "short": "Sun"},
]
DAY_KEYS = [d["key"] for d in DAYS]

STATUS_ORDER = ["running", "technical", "major"]
STATUS = {
    "running":   {"name": "🟢 Running",               "hex": "#10b981"},
    "technical": {"name": "🟣 Technical intervention","hex": "#8b5cf6"},
    "major":     {"name": "🔴 Major component (GC)",  "hex": "#ef4444"},
}
STATUS_NAMES = [STATUS[k]["name"] for k in STATUS_ORDER]
NAME_TO_KEY = {v["name"]: k for k, v in STATUS.items()}
SEVERITY = {"running": 0, "technical": 1, "major": 2}

INTERVALS = [
    {"h": h, "label": f"{h:02d}:00–{(h+1)%24:02d}:00"}
    for h in range(24)
]
SLOT_HOURS = 1
DEFAULT_INTERVALS = ["running"] * 24

DEFAULT_ROSTER = {
    "G52": [f"{i:02d}" for i in range(1, 13)],
    "G80": [f"{i:02d}" for i in range(1, 12)],
}

SEED_DAY_MACHINES = {
    "G52": {
        **{f"{i:02d}": {"intervals": DEFAULT_INTERVALS[:], "notes": ""} for i in range(1, 13)},
        "07": {"intervals": ["technical"] * 24, "notes": "Cable delivery pending (end of August)"},
        "11": {"intervals": ["technical"] * 24, "notes": "Wind vane issue"},
    },
    "G80": {
        **{f"{i:02d}": {"intervals": DEFAULT_INTERVALS[:], "notes": ""} for i in range(1, 12)},
        "01": {"intervals": ["major"] * 24, "notes": "GC (Blade bearing)"},
        "02": {"intervals": ["major"] * 24, "notes": "GC (Blade bearing)"},
        "03": {"intervals": ["technical"] * 24, "notes": "Gear motors"},
        "04": {"intervals": ["technical"] * 24, "notes": "Under inspection"},
        "08": {"intervals": ["technical"] * 24, "notes": "Yaw system"},
        "09": {"intervals": ["major"] * 24, "notes": "GC (Generator)"},
        "10": {"intervals": ["technical"] * 24, "notes": "Hub power supply"},
    },
}
SEED_PLAN = {"G52": "11", "G80": "04, 08"}

DEFAULT_GC = [
    {"id": "gc1", "group": "G80", "machine": "09", "component": "Generator", "due": "End of August"},
    {"id": "gc2", "group": "G80", "machine": "01", "component": "Blade bearing", "due": "End of September"},
    {"id": "gc3", "group": "G80", "machine": "02", "component": "Blade bearing", "due": "End of September"},
]

EN_MONTHS = ["January", "February", "March", "April", "May", "June", "July", "August",
             "September", "October", "November", "December"]
DB_PATH = "wind_farm.db"

# ============================================================
# Persistence
# ============================================================

def get_conn():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.execute("CREATE TABLE IF NOT EXISTS kv_store (key TEXT PRIMARY KEY, value TEXT, updated_at TEXT)")
    return conn


def kv_get(key, default=None):
    conn = get_conn()
    try:
        row = conn.execute("SELECT value FROM kv_store WHERE key=?", (key,)).fetchone()
    finally:
        conn.close()
    if row is None:
        return default
    try:
        return json.loads(row[0])
    except (json.JSONDecodeError, TypeError):
        return default


def kv_set(key, value):
    conn = get_conn()
    try:
        conn.execute(
            "INSERT INTO kv_store (key, value, updated_at) VALUES (?, ?, ?) "
            "ON CONFLICT(key) DO UPDATE SET value=excluded.value, updated_at=excluded.updated_at",
            (key, json.dumps(value), datetime.now().isoformat()),
        )
        conn.commit()
    finally:
        conn.close()


def save_all():
    kv_set("roster", st.session_state.roster)
    kv_set("week_data", st.session_state.week_data)
    kv_set("gc_planning", st.session_state.gc_planning)
    kv_set("settings", st.session_state.settings)
    kv_set("history", st.session_state.history)

# ============================================================
# Data helpers
# ============================================================

def empty_machines(roster):
    return {g: {mid: {"intervals": DEFAULT_INTERVALS[:], "notes": ""} for mid in ids} for g, ids in roster.items()}


def empty_day(roster):
    return {"machines": empty_machines(roster), "plan": {"G52": "", "G80": ""}, "production_mwh": 0.0}


def today_key():
    return DAY_KEYS[date.today().weekday()]


def build_default_week():
    tk = today_key()
    week = {}
    for d in DAYS:
        if d["key"] == tk:
            week[d["key"]] = {"machines": json.loads(json.dumps(SEED_DAY_MACHINES)), "plan": dict(SEED_PLAN)}
        else:
            week[d["key"]] = empty_day(DEFAULT_ROSTER)
    return week


def empty_week_for_roster(roster):
    return {d["key"]: empty_day(roster) for d in DAYS}


def normalize_machine(m):
    if not m:
        return {"intervals": DEFAULT_INTERVALS[:], "notes": ""}
    intervals = m.get("intervals", [])
    if isinstance(intervals, list) and len(intervals) == 6:
        expanded = []
        for s in intervals:
            s = "technical" if s == "waiting" else s
            s = s if s in STATUS_ORDER else "running"
            expanded.extend([s] * 4)
        intervals = expanded
    elif isinstance(intervals, list) and len(intervals) == 24:
        intervals = ["technical" if s == "waiting" else (s if s in STATUS_ORDER else "running") for s in intervals]
    else:
        intervals = DEFAULT_INTERVALS[:]
    return {"intervals": intervals, "notes": m.get("notes", m.get("constat", ""))}


def normalize_week_data(week_data):
    if not week_data:
        return build_default_week()
    result = {}
    for d in DAYS:
        day = week_data.get(d["key"], {})
        machines = {}
        for group, ids_map in day.get("machines", {}).items():
            machines[group] = {mid: normalize_machine(mv) for mid, mv in ids_map.items()}
        result[d["key"]] = {"machines": machines, "plan": day.get("plan", {"G52": "", "G80": ""}), "production_mwh": float(day.get("production_mwh", 0.0) or 0.0)}
    return result


def week_range(ref=None):
    ref = ref or date.today()
    monday = ref - timedelta(days=ref.weekday())
    return monday, monday + timedelta(days=6)


def iso_date(d):
    return d.strftime("%Y-%m-%d")


def parse_iso_date(s):
    return datetime.strptime(s, "%Y-%m-%d").date()


def format_date(d):
    return f"{d.day} {EN_MONTHS[d.month - 1]}"


def default_settings():
    return {"active_week_key": iso_date(week_range()[0])}


def day_worst_status(intervals):
    worst = "running"
    for s in intervals or []:
        if SEVERITY.get(s, 0) > SEVERITY.get(worst, 0):
            worst = s
    return worst


def machine_week_stats(group, mid, week_data):
    counts = {k: 0 for k in STATUS_ORDER}
    for d in DAYS:
        m = week_data.get(d["key"], {}).get("machines", {}).get(group, {}).get(mid, {})
        for s in normalize_machine(m)["intervals"]:
            counts[s] = counts.get(s, 0) + 1
    total_slots = sum(counts.values())
    running_slots = counts["running"]
    availability = (running_slots / total_slots * 100) if total_slots else 100.0
    downtime_h = (total_slots - running_slots) * SLOT_HOURS
    return {"counts": counts, "availability_pct": availability, "downtime_h": downtime_h}


def day_group_availability(group, day_key, roster, week_data):
    ids = roster.get(group, [])
    total = len(ids) * len(INTERVALS)
    running = 0
    for mid in ids:
        m = week_data.get(day_key, {}).get("machines", {}).get(group, {}).get(mid, {})
        running += sum(1 for s in normalize_machine(m)["intervals"] if s == "running")
    return (running / total * 100) if total else 100.0


def compute_performance(roster, week_data):
    group_stats = {}
    status_slots_total = {k: 0 for k in STATUS_ORDER}
    total_slots_all = 0
    running_slots_all = 0

    for group, ids in roster.items():
        machine_rows = []
        group_counts = {k: 0 for k in STATUS_ORDER}
        for mid in ids:
            mstats = machine_week_stats(group, mid, week_data)
            for k in STATUS_ORDER:
                group_counts[k] += mstats["counts"][k]
                status_slots_total[k] += mstats["counts"][k]
            machine_rows.append({
                "machine": mid,
                "availability_pct": mstats["availability_pct"],
                "downtime_h": mstats["downtime_h"],
                "technical_h": mstats["counts"]["technical"] * SLOT_HOURS,
                "major_h": mstats["counts"]["major"] * SLOT_HOURS,
            })

        group_total = sum(group_counts.values())
        group_running = group_counts["running"]
        group_avail = (group_running / group_total * 100) if group_total else 100.0
        group_stats[group] = {
            "availability_pct": group_avail,
            "downtime_h": (group_total - group_running) * SLOT_HOURS,
            "status_counts": group_counts,
            "machines": machine_rows,
            "daily_availability": {d["key"]: day_group_availability(group, d["key"], roster, week_data) for d in DAYS},
        }
        total_slots_all += group_total
        running_slots_all += group_running

    overall_avail = (running_slots_all / total_slots_all * 100) if total_slots_all else 100.0
    return {
        "overall_availability_pct": overall_avail,
        "overall_downtime_h": (total_slots_all - running_slots_all) * SLOT_HOURS,
        "status_counts": status_slots_total,
        "group_stats": group_stats,
    }


def compute_report_data(roster, week_data, gc_planning):
    perf = compute_performance(roster, week_data)
    daily_plan = [{
        "label": d["label"],
        "G52": week_data.get(d["key"], {}).get("plan", {}).get("G52", ""),
        "G80": week_data.get(d["key"], {}).get("plan", {}).get("G80", ""),
    } for d in DAYS]
    production = [{"label": d["label"], "value": float(week_data.get(d["key"], {}).get("production_mwh", 0.0) or 0.0)} for d in DAYS]
    return {"performance": perf, "daily_plan": daily_plan, "gc_planning": gc_planning, "production": production}

# ============================================================
# Charts
# ============================================================

def make_gauge(value, title_text):
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=value,
        number={"suffix": "%", "font": {"size": 26}},
        title={"text": title_text, "font": {"size": 13}},
        gauge={"axis": {"range": [0, 100]}, "bar": {"color": "#0f766e", "thickness": 0.28}},
    ))
    fig.update_layout(height=190, margin=dict(l=25, r=25, t=45, b=10))
    return fig


def make_real_production_chart(week_data):
    values = [float(week_data.get(d["key"], {}).get("production_mwh", 0.0) or 0.0) for d in DAYS]
    fig = go.Figure(go.Bar(
        x=[d["short"] for d in DAYS], y=values,
        text=[f"{v:.1f}" for v in values], textposition="auto",
        name="Real production"
    ))
    fig.update_layout(
        title="Real daily production — whole wind farm",
        yaxis_title="Production (MWh/day)", xaxis_title="Day",
        height=340, margin=dict(l=40, r=20, t=50, b=40)
    )
    return fig


def make_status_hours_chart(perf):
    labels = [STATUS[k]["name"].split(" ", 1)[1] for k in ["technical", "major"]]
    hours = [perf["status_counts"][k] * SLOT_HOURS for k in ["technical", "major"]]
    fig = go.Figure(go.Bar(x=labels, y=hours))
    fig.update_layout(yaxis_title="Hours", height=320, margin=dict(l=40, r=20, t=20, b=60))
    return fig

# ============================================================
# Reports
# ============================================================

def build_report_docx(data, monday, sunday):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)
    teal = RGBColor(0x13, 0x4e, 0x4a)

    title = doc.add_heading("Weekly Performance Report — Wind Farm G52 & G80", level=1)
    title.runs[0].font.color.rgb = teal
    doc.add_paragraph(f"Week of {format_date(monday)} to {format_date(sunday)}")

    perf = data["performance"]
    t = doc.add_table(rows=2, cols=4)
    t.style = "Light Grid Accent 1"
    vals = [
        f"{perf['overall_availability_pct']:.1f}%",
        f"{perf['group_stats'].get('G52', {}).get('availability_pct', 100):.1f}%",
        f"{perf['group_stats'].get('G80', {}).get('availability_pct', 100):.1f}%",
        f"{perf['overall_downtime_h']:.0f} h",
    ]
    labels = ["Overall availability", "G52 availability", "G80 availability", "Total downtime"]
    for i in range(4):
        t.cell(0, i).paragraphs[0].add_run(vals[i]).bold = True
        t.cell(1, i).text = labels[i]

    doc.add_heading("Real production — whole wind farm", level=2)
    pt = doc.add_table(rows=1, cols=2)
    pt.style = "Light List Accent 1"
    pt.rows[0].cells[0].text = "Day"
    pt.rows[0].cells[1].text = "Production (MWh)"
    for r in data.get("production", []):
        row = pt.add_row().cells
        row[0].text = r["label"]
        row[1].text = f'{r["value"]:.1f}'

    doc.add_heading("Machine performance", level=2)
    mt = doc.add_table(rows=1, cols=6)
    mt.style = "Light List Accent 1"
    headers = ["Group", "Machine", "Availability", "Downtime", "Technical", "Major GC"]
    for i, h in enumerate(headers):
        mt.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for group, stats in perf["group_stats"].items():
        for r in stats["machines"]:
            row = mt.add_row().cells
            vals = [group, r["machine"], f"{r['availability_pct']:.1f}%", f"{r['downtime_h']:.0f} h",
                    f"{r['technical_h']:.0f} h", f"{r['major_h']:.0f} h"]
            for i, v in enumerate(vals):
                row[i].text = str(v)

    doc.add_heading("Intervention log", level=2)
    plans = [r for r in data["daily_plan"] if r["G52"] or r["G80"]]
    if not plans:
        doc.add_paragraph("No interventions planned this week.")
    for r in plans:
        parts = []
        if r["G52"]: parts.append(f"G52 ({r['G52']})")
        if r["G80"]: parts.append(f"G80 ({r['G80']})")
        doc.add_paragraph(f"{r['label']}: {' · '.join(parts)}", style="List Bullet")

    doc.add_heading("GC Planning — major components", level=2)
    if not data["gc_planning"]:
        doc.add_paragraph("No major-component deadlines scheduled.")
    else:
        gt = doc.add_table(rows=1, cols=4)
        gt.style = "Light List Accent 1"
        for i, h in enumerate(["Group", "Machine", "Component", "Due date"]):
            gt.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
        for r in data["gc_planning"]:
            row = gt.add_row().cells
            for i, v in enumerate([r["group"], r["machine"], r["component"], r["due"]]):
                row[i].text = v

    doc.add_paragraph(f"Generated on {datetime.now().strftime('%d %B %Y, %H:%M')}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ============================================================
# Session state and rollover
# ============================================================

def init_state():
    if "initialized" in st.session_state:
        return
    st.session_state.roster = kv_get("roster", DEFAULT_ROSTER)
    st.session_state.week_data = normalize_week_data(kv_get("week_data", None))
    st.session_state.gc_planning = kv_get("gc_planning", DEFAULT_GC)
    st.session_state.settings = kv_get("settings", default_settings())
    st.session_state.history = kv_get("history", [])
    st.session_state.initialized = True
    check_week_rollover()


def archive_week(week_key, roster, week_data, gc_planning):
    entry = {
        "week_key": week_key,
        "archived_at": datetime.now().isoformat(),
        "roster": roster,
        "week_data": week_data,
        "gc_planning": gc_planning,
    }
    st.session_state.history = [entry] + [h for h in st.session_state.history if h.get("week_key") != week_key]
    st.session_state.history.sort(key=lambda h: h["week_key"], reverse=True)


def check_week_rollover():
    real_key = iso_date(week_range()[0])
    active_key = st.session_state.settings.get("active_week_key")
    if active_key and active_key != real_key:
        archive_week(active_key, st.session_state.roster, st.session_state.week_data, st.session_state.gc_planning)
        st.session_state.week_data = empty_week_for_roster(st.session_state.roster)
        st.session_state.settings["active_week_key"] = real_key
        save_all()


def close_week():
    week_key = st.session_state.settings.get("active_week_key") or iso_date(week_range()[0])
    archive_week(week_key, st.session_state.roster, st.session_state.week_data, st.session_state.gc_planning)
    st.session_state.week_data = empty_week_for_roster(st.session_state.roster)
    st.session_state.settings["active_week_key"] = iso_date(week_range()[0])
    save_all()

# ============================================================
# UI
# ============================================================

def render_machine_table(group, day_key):
    roster_ids = st.session_state.roster[group]
    machines = st.session_state.week_data[day_key]["machines"].setdefault(group, {})
    pct = day_group_availability(group, day_key, st.session_state.roster, st.session_state.week_data)
    st.markdown(f"### {group} — {pct:.1f}% available")
    st.caption("24 hourly controls per turbine — arranged 8 + 8 + 8.")

    for mid in roster_ids:
        m = machines.setdefault(mid, {"intervals": DEFAULT_INTERVALS[:], "notes": ""})
        with st.expander(f"{group} — Turbine {mid}", expanded=False):
            intervals = normalize_machine(m)["intervals"]
            for start in (0, 8, 16):
                cols = st.columns(8)
                for j, hour in enumerate(range(start, start + 8)):
                    with cols[j]:
                        current = intervals[hour] if hour < len(intervals) else "running"
                        selected = st.selectbox(
                            f"{hour:02d}:00", STATUS_NAMES,
                            index=STATUS_NAMES.index(STATUS[current]["name"]),
                            key=f"hour_{group}_{day_key}_{mid}_{hour}"
                        )
                        intervals[hour] = NAME_TO_KEY[selected]
            m["intervals"] = intervals
            m["notes"] = st.text_input("Notes", value=m.get("notes", ""), key=f"notes_{group}_{day_key}_{mid}")


def render_gc_planning():
    st.markdown("##### GC Planning — major components")
    rows = st.session_state.gc_planning
    df = pd.DataFrame(rows)[["group", "machine", "component", "due"]] if rows else pd.DataFrame(columns=["group", "machine", "component", "due"])
    df.columns = ["Group", "Machine", "Component", "Due date"]
    edited = st.data_editor(df, hide_index=True, use_container_width=True, num_rows="dynamic", key="gc_editor",
                            column_config={"Group": st.column_config.SelectboxColumn(options=["G52", "G80"])})
    new_rows = []
    for i, r in edited.iterrows():
        rid = rows[i]["id"] if i < len(rows) else f"gc{datetime.now().timestamp()}_{i}"
        new_rows.append({"id": rid, "group": r["Group"] or "G52", "machine": r["Machine"] or "",
                         "component": r["Component"] or "", "due": r["Due date"] or ""})
    st.session_state.gc_planning = new_rows


def render_day_view():
    tk = today_key()
    sel = st.radio("Day", [d["short"] for d in DAYS], index=DAY_KEYS.index(tk), horizontal=True, label_visibility="collapsed")
    day_key = DAY_KEYS[[d["short"] for d in DAYS].index(sel)]

    current_prod = float(st.session_state.week_data[day_key].get("production_mwh", 0.0) or 0.0)
    st.session_state.week_data[day_key]["production_mwh"] = st.number_input(
        "⚡ Real production of the whole wind farm (G52 + G80) — MWh/day",
        min_value=0.0, value=current_prod, step=1.0, key=f"production_{day_key}"
    )

    g52_tab, g80_tab = st.tabs(["G52 turbines", "G80 turbines"])
    with g52_tab:
        render_machine_table("G52", day_key)
    with g80_tab:
        render_machine_table("G80", day_key)

    st.markdown("##### Today's intervention plan")
    plan = st.session_state.week_data[day_key]["plan"]
    p1, p2 = st.columns(2)
    with p1: plan["G52"] = st.text_input("G52", value=plan.get("G52", ""), key=f"plan52_{day_key}")
    with p2: plan["G80"] = st.text_input("G80", value=plan.get("G80", ""), key=f"plan80_{day_key}")
    render_gc_planning()
    save_all()


def render_week_view():
    st.markdown("### Weekly availability (%)")
    for group in ["G52", "G80"]:
        st.markdown(f"### {group}")
        rows = []
        for mid in st.session_state.roster[group]:
            row = {"Machine": mid}
            for d in DAYS:
                m = st.session_state.week_data[d["key"]]["machines"].get(group, {}).get(mid, {"intervals": DEFAULT_INTERVALS})
                intervals = normalize_machine(m)["intervals"]
                row[d["short"]] = sum(1 for s in intervals if s == "running") / 24 * 100
            rows.append(row)
        df = pd.DataFrame(rows).set_index("Machine")
        st.dataframe(df.style.format("{:.1f}%"), use_container_width=True)
        st.markdown("---")


def render_performance_view(roster=None, week_data=None, key_prefix="performance"):
    roster = roster or st.session_state.roster
    week_data = week_data or st.session_state.week_data
    perf = compute_performance(roster, week_data)

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Overall availability", f"{perf['overall_availability_pct']:.1f}%")
    c2.metric("G52 availability", f"{perf['group_stats'].get('G52', {}).get('availability_pct', 100):.1f}%")
    c3.metric("G80 availability", f"{perf['group_stats'].get('G80', {}).get('availability_pct', 100):.1f}%")
    c4.metric("Weekly real production", f"{sum(float(week_data.get(d['key'], {}).get('production_mwh', 0.0) or 0.0) for d in DAYS):.1f} MWh")

    st.plotly_chart(make_real_production_chart(week_data), use_container_width=True, key=f"{key_prefix}_real_production")

    left, right = st.columns(2)
    with left:
        st.markdown("##### Downtime by status")
        st.plotly_chart(make_status_hours_chart(perf), use_container_width=True, key=f"{key_prefix}_status_hours")
    with right:
        st.markdown("##### Group downtime")
        for group, stats in perf["group_stats"].items():
            st.metric(group, f"{stats['downtime_h']:.0f} h", f"Availability {stats['availability_pct']:.1f}%")

    st.markdown("##### Turbine performance ranking")
    rows = []
    for group, stats in perf["group_stats"].items():
        for r in stats["machines"]:
            rows.append({
                "Group": group,
                "Machine": r["machine"],
                "Availability (%)": round(r["availability_pct"], 1),
                "Downtime (h)": round(r["downtime_h"], 1),
                "Technical (h)": round(r["technical_h"], 1),
                "Major GC (h)": round(r["major_h"], 1),
            })
    ranking = pd.DataFrame(rows).sort_values(["Availability (%)", "Downtime (h)"], ascending=[True, False])
    st.dataframe(ranking, hide_index=True, use_container_width=True)
    st.caption("Lower availability / higher downtime appears first, making the most critical turbines easy to identify.")


def render_report_view():
    monday, sunday = week_range()
    data = compute_report_data(st.session_state.roster, st.session_state.week_data, st.session_state.gc_planning)
    st.markdown(f"### Weekly Performance Report — {format_date(monday)} to {format_date(sunday)}")
    docx = build_report_docx(data, monday, sunday)
    st.download_button("⬇️ Download Word (.docx)", data=docx, file_name=f"performance-report-{iso_date(monday)}.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    render_performance_view(key_prefix="report")

    st.markdown("##### Daily real production")
    st.dataframe(pd.DataFrame([{"Day": r["label"], "Production (MWh)": r["value"]} for r in data.get("production", [])]), hide_index=True, use_container_width=True)

    st.markdown("##### Intervention log")
    plans = [r for r in data["daily_plan"] if r["G52"] or r["G80"]]
    if not plans:
        st.caption("No interventions planned this week.")
    for r in plans:
        parts = []
        if r["G52"]: parts.append(f"G52 ({r['G52']})")
        if r["G80"]: parts.append(f"G80 ({r['G80']})")
        st.markdown(f"**{r['label']}:** {' · '.join(parts)}")

    st.markdown("##### GC Planning — major components")
    if data["gc_planning"]:
        st.dataframe(pd.DataFrame(data["gc_planning"])[["group", "machine", "component", "due"]], hide_index=True, use_container_width=True)
    else:
        st.caption("No major-component deadlines scheduled.")


def render_history_view():
    history = st.session_state.history
    if not history:
        st.caption("No archived weeks yet.")
        return
    labels = []
    for entry in history:
        monday = parse_iso_date(entry["week_key"])
        labels.append(f"Week of {format_date(monday)} to {format_date(monday + timedelta(days=6))}")
    idx = st.selectbox("Select archived week", options=range(len(history)), format_func=lambda i: labels[i])
    entry = history[idx]
    render_performance_view(entry["roster"], entry["week_data"], key_prefix=f"history_{idx}")

# ============================================================
# Main
# ============================================================

def main():
    st.set_page_config(page_title="Wind Farm Performance G52 & G80", page_icon="🌬️", layout="wide")
    init_state()

    with st.sidebar:
        st.markdown("## 🌬️ Wind Farm Performance")
        st.caption("G52 & G80 — Hourly technical availability & real production study")
        st.write(datetime.now().strftime("%A %d %B %Y — %H:%M"))
        st.markdown("---")
        if st.button("📦 New week", use_container_width=True, help="Archive current week and start a new one"):
            close_week()
            st.rerun()
        st.caption("No forecast module: this app focuses on technical performance, hourly availability and real production.")

    tabs = st.tabs(["📅 Day view", "🗓️ Week view", "📊 Performance", "📄 Report", "📚 History"])
    with tabs[0]: render_day_view()
    with tabs[1]: render_week_view()
    with tabs[2]: render_performance_view(key_prefix="performance_tab")
    with tabs[3]: render_report_view()
    with tabs[4]: render_history_view()


if __name__ == "__main__":
    main()
